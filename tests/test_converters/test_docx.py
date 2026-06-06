import tempfile
from pathlib import Path
import docx
from docx.shared import Pt
from core.converters.docx import convert_docx

def _make_docx(tmp_dir: Path) -> Path:
    doc = docx.Document()
    doc.add_heading("Main Title", level=1)
    doc.add_heading("Subtitle", level=2)
    p = doc.add_paragraph()
    run = p.add_run("bold text")
    run.bold = True
    p.add_run(" and ")
    run2 = p.add_run("italic text")
    run2.italic = True
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Val 1"
    table.cell(1, 1).text = "Val 2"
    path = tmp_dir / "sample.docx"
    doc.save(str(path))
    return path

def test_docx_headings():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        docx_path = _make_docx(tmp)
        md, assets = convert_docx(docx_path, tmp / "assets")
        assert "# Main Title" in md
        assert "## Subtitle" in md

def test_docx_inline_formatting():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        docx_path = _make_docx(tmp)
        md, assets = convert_docx(docx_path, tmp / "assets")
        assert "**bold text**" in md
        assert "_italic text_" in md

def test_docx_table():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        docx_path = _make_docx(tmp)
        md, assets = convert_docx(docx_path, tmp / "assets")
        assert "| Header A | Header B |" in md
        assert "| Val 1 | Val 2 |" in md

def test_docx_returns_asset_list():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        docx_path = _make_docx(tmp)
        md, assets = convert_docx(docx_path, tmp / "assets")
        assert isinstance(assets, list)
